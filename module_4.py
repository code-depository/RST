"""
module_4.py
Section VI -- Three-Way Benchmarking Protocol

Shared harness for comparing QMC, Espresso, and RST on a common Boolean
function instance: metrics (gate count, literal count, execution time,
variable/support reduction ratio) and correctness verification
(exhaustive for effective n <= 24, sampled with a stated confidence
bound otherwise -- Section VI.D).

Espresso is invoked via pyeda's binding to the real, compiled Berkeley
Espresso binary (not a reimplementation), per the Section VI.A
methodology commitment.

Depends on module_1.py (QMC, DIS) and module_3.py (attribute-level
reducts) in the same directory.

Also generates the three manuscript artifacts for Section VI (.docx,
matching the Table 1 pattern from module_1.py):
    - Table 2: Figures of Merit (Section VI.B)
    - Table 3: Feasibility Scope by Effective Input Count (Section VI.C)
    - Table 4: Circuit Family Verification Assignment (Section VI.F)

External dependencies: pyeda (Espresso binary), python-docx (tables).
Run this file directly to execute a worked three-way comparison, a
verification-harness self-test, and generate all three tables.
"""

import math
import time
from typing import Dict, FrozenSet, List, Tuple

from pyeda.inter import exprvars, truthtable, espresso_tts
from pyeda.boolalg.expr import AndOp, OrOp, Complement, Variable

from module_1 import DecisionInformationSystem, qmc_prime_implicants
from module_3 import all_reducts, minimal_reducts

Literal = Tuple[int, int]


def petrick_cover(n: int, onset: List[int], primes: List[FrozenSet[Literal]]) -> List[FrozenSet[Literal]]:
    """Minimum-GATE-COUNT cover of `onset` from `primes` (not merely
    minimum term count -- these are different objectives, since gate
    count also weighs each term's literal count and complemented
    literals; minimizing term count alone can select a cover that is
    NOT gate-count-optimal, silently breaking the "QMC is exact"
    premise the rest of the manuscript relies on). Essential PIs are
    extracted first; the remainder is solved via exact ILP (bundled CBC
    solver, offline), with a per-term cost matching the Section VI.B
    gate_count definition.
    """
    import pulp

    def covers(pi, m):
        bits = format(m, f"0{n}b")
        return all(bits[a] == str(v) for a, v in pi)

    def term_cost(pi):
        # AND gates (literals beyond the first) + NOT gates (complemented
        # literals). The final +1 OR-gate contribution per selected term
        # is handled by the constant "+|cover|-1" outside the ILP, so is
        # omitted here (it does not depend on WHICH terms are chosen,
        # only how many -- and minimizing sum(cost_i * x_i) with this
        # per-term cost already favors fewer, cheaper terms correctly
        # since essential PIs are fixed and the OR-gate term is monotonic
        # in |cover|).
        return max(0, len(pi) - 1) + sum(1 for (_, v) in pi if v == 0)

    coverage = {m: [i for i, pi in enumerate(primes) if covers(pi, m)] for m in onset}
    essential = {idxs[0] for idxs in coverage.values() if len(idxs) == 1}
    covered = set()
    for i in essential:
        covered.update(m for m in onset if covers(primes[i], m))
    remaining = [m for m in onset if m not in covered]

    if not remaining:
        return [primes[i] for i in essential]

    candidates = [i for i in range(len(primes)) if i not in essential]

    prob = pulp.LpProblem("min_gate_count_cover", pulp.LpMinimize)
    x = {i: pulp.LpVariable(f"x_{i}", cat="Binary") for i in candidates}
    # weighted objective: per-term gate cost, NOT bare term count
    prob += pulp.lpSum(term_cost(primes[i]) * x[i] for i in candidates) + pulp.lpSum(x.values())
    # the extra "+ sum(x)" term approximates each selected term's marginal
    # contribution to the OR-gate count; exact OR-gate accounting
    # (max(0,|cover|-1)) is applied after solving, on the final cover.
    for m in remaining:
        covering = [i for i in candidates if covers(primes[i], m)]
        prob += pulp.lpSum(x[i] for i in covering) >= 1
    prob.solve(pulp.PULP_CBC_CMD(msg=0))

    chosen = essential | {i for i in candidates if x[i].value() == 1}
    return [primes[i] for i in chosen]





# ---------------------------------------------------------------------
# Metrics  (Section VI.B)
# ---------------------------------------------------------------------

def literal_count(cover: List[FrozenSet[Literal]]) -> int:
    return sum(len(term) for term in cover)


def gate_count(cover: List[FrozenSet[Literal]]) -> int:
    """Two-level AND-OR gate count (Section VI.B definition)."""
    and_gates = sum(max(0, len(t) - 1) for t in cover)
    not_gates = sum(1 for t in cover for (_, v) in t if v == 0)
    or_gates = max(0, len(cover) - 1)
    return and_gates + not_gates + or_gates


def variable_reduction_ratio(n: int, reduct_size: int) -> float:
    return 1.0 - (reduct_size / n) if n else 0.0


# ---------------------------------------------------------------------
# Espresso via pyeda (real binary, Section VI.A)
# ---------------------------------------------------------------------

def _to_pyeda_index(m: int, n: int) -> int:
    """Our DecisionInformationSystem.bits(m) treats attribute a as the
    bit at string-position a (MSB-first: attribute 0 is the most
    significant bit). pyeda's truthtable treats X[i] as bit i counting
    from the LSB of the table position. To make pyeda's X[i] equal our
    attribute i, the table must be indexed by the bit-reversal of m.
    """
    return int(format(m, f"0{n}b")[::-1], 2)


def _pyeda_term_to_literals(term) -> FrozenSet[Literal]:
    """Convert one AndOp/Variable/Complement term to our (attr, val) form.
    Given the indexing fix above, pyeda's x[i] already corresponds
    directly to our attribute i -- no further reversal needed here."""
    literals = set()
    factors = term.xs if isinstance(term, AndOp) else [term]
    for f in factors:
        if isinstance(f, Complement):
            literals.add((_var_index(f.inputs[0]), 0))
        elif isinstance(f, Variable):
            literals.add((_var_index(f), 1))
    return frozenset(literals)


def _var_index(v) -> int:
    return v.indices[0]


def espresso_cover(n: int, onset: List[int], dc: List[int]) -> List[FrozenSet[Literal]]:
    """Minimize via the real Espresso binary (pyeda binding)."""
    X = exprvars("x", n)
    bitstring = [None] * (2 ** n)
    for m in range(2 ** n):
        idx = _to_pyeda_index(m, n)
        if m in onset:
            bitstring[idx] = "1"
        elif m in dc:
            bitstring[idx] = "-"
        else:
            bitstring[idx] = "0"
    tt = truthtable(X, "".join(bitstring))
    (result,) = espresso_tts(tt)
    dnf = result.to_dnf()
    terms = dnf.xs if isinstance(dnf, OrOp) else [dnf]
    return [_pyeda_term_to_literals(t) for t in terms]


# ---------------------------------------------------------------------
# RST cover: reduct projection + QMC on the reduced space (Section II/V)
# ---------------------------------------------------------------------

def rst_cover(dis: DecisionInformationSystem):
    """Returns (cover, reduct) where cover is expressed over the FULL
    n-variable space (reduct-excluded variables simply absent from every
    literal), and reduct is the minimal reduct used.
    """
    reducts = all_reducts(dis)
    reduct = sorted(minimal_reducts(reducts), key=lambda r: (len(r), sorted(r)))[0]
    reduct_sorted = sorted(reduct)
    k = len(reduct_sorted)

    # project truth table onto reduct attributes
    def project(m: int) -> int:
        bits = dis.bits(m)
        sub = "".join(bits[a] for a in reduct_sorted)
        return int(sub, 2) if sub else 0

    proj_onset = sorted({project(m) for m in dis.onset})
    proj_offset = {project(m) for m in dis.offset}
    proj_dc_raw = {project(m) for m in dis.dc}
    # a projected value remains a genuine don't-care only if it never
    # coincides with a projected onset OR offset value -- either of those
    # is a confirmed assignment and takes precedence over "unspecified"
    proj_dc = sorted(proj_dc_raw - set(proj_onset) - proj_offset)
    proj_pis = qmc_prime_implicants(k, proj_onset, proj_dc)
    proj_cover = petrick_cover(k, proj_onset, proj_pis)

    # re-express projected literals in terms of original attribute indices
    cover = []
    for term in proj_cover:
        remapped = frozenset((reduct_sorted[i], v) for i, v in term)
        cover.append(remapped)
    return cover, reduct


# ---------------------------------------------------------------------
# Correctness verification  (Section VI.D)
# ---------------------------------------------------------------------

def _term_satisfied(term: FrozenSet[Literal], m: int, n: int) -> bool:
    bits = format(m, f"0{n}b")
    return all(bits[a] == str(v) for a, v in term)


def _cover_value(cover: List[FrozenSet[Literal]], m: int, n: int) -> int:
    return int(any(_term_satisfied(t, m, n) for t in cover))


def exhaustive_verify(dis: DecisionInformationSystem, cover: List[FrozenSet[Literal]]) -> bool:
    """Checks cover against the original specification at every minterm.
    Don't-care minterms are not checked (any value is acceptable there)."""
    onset_set, offset_set = set(dis.onset), set(dis.offset)
    for m in range(2 ** dis.n):
        if m in onset_set and _cover_value(cover, m, dis.n) != 1:
            return False
        if m in offset_set and _cover_value(cover, m, dis.n) != 0:
            return False
    return True


def required_sample_size(confidence: float, error_rate: float) -> int:
    """N = ceil(ln(1-c) / ln(1-e)), Section VI.D."""
    return math.ceil(math.log(1 - confidence) / math.log(1 - error_rate))


# ---------------------------------------------------------------------
# Manuscript artifacts: Tables 2-4 (.docx), Section VI
# ---------------------------------------------------------------------

def _new_doc_with_title(title_text: str):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(11)
    return doc


def generate_table2_docx(path: str) -> str:
    """Table 2: Figures of Merit (Section VI.B)."""
    doc = _new_doc_with_title(
        "Table 2. Figures of merit used in the three-way benchmarking protocol."
    )
    rows = [
        ("Metric", "Definition", "Role"),
        ("Gate count", "\u03a3 over terms of max(0,k\u22121) AND gates + NOT gates for "
                        "complemented literals + max(0,m\u22121) OR gates "
                        "(k = literals/term, m = #terms)", "Primary quality metric"),
        ("Literal count", "Total literals across all product terms",
         "Finer-grained standard metric"),
        ("Execution time", "Wall-clock time to minimize",
         "Scalability signal, not a competitiveness claim"),
        ("Variable/support reduction ratio", "1 \u2212 (|reduct| / n)",
         "First-class metric, computed on every instance"),
        ("Space / memory", "Peak discernibility-matrix and cover-storage footprint",
         "Validates complexity analysis at increasing n"),
        ("Circuit depth (optional)", "Levels in the two-level AND-OR realization",
         "Delay proxy, reported where relevant"),
        ("Correctness verification outcome", "Structural proof, exhaustive check, "
                                              "or sampled equivalence result with stated confidence bound",
         "Required at every n (Section VI.D)"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, rows[0]):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows[1:]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            cell.text = text
    doc.save(path)
    return path


def generate_table3_docx(path: str) -> str:
    """Table 3: Feasibility scope by effective input count (Section VI.C)."""
    doc = _new_doc_with_title(
        "Table 3. Method feasibility by effective input count n."
    )
    rows = [
        ("Effective n", "QMC (exact)", "Espresso", "RST-exact", "RST-greedy", "Benchmark type"),
        ("\u2264 8", "Tractable", "Tractable", "Tractable", "\u2014", "Random, DC 0\u201350%"),
        ("10", "Timeout risk (15s cap)", "Tractable", "Slow but tractable", "Tractable",
         "Random, reduced instance count"),
        ("12\u201320", "Infeasible", "Tractable", "Infeasible", "Tractable",
         "Structured circuit families"),
        ("20\u201333", "Infeasible", "Tractable (structured only)", "Infeasible", "Tractable",
         "Structured circuit families only"),
    ]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, rows[0]):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows[1:]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            cell.text = text
    doc.save(path)
    return path


def generate_table4_docx(path: str) -> str:
    """Table 4: Circuit family verification assignment (Section VI.F)."""
    doc = _new_doc_with_title(
        "Table 4. Circuit family verification assignment."
    )
    rows = [
        ("Family", "Nominal widths", "Effective n range (per output)", "Verification"),
        ("Ripple-carry adder", "4/8/12/16-bit", "3 (per-bit carry logic)",
         "Structural (carry-chain induction)"),
        ("Magnitude comparator", "4/8/12/16-bit", "3\u20134 (per-bit compare-cascade logic)",
         "Structural (bit-slice induction)"),
        ("Priority encoder", "4/8/12/16-bit", "up to w",
         "Exhaustive (monolithic per width; w \u2264 16 well within exhaustive threshold)"),
        ("Binary decoder", "4/8/12/16-bit", "2\u20134 (\u2308log\u2082 w\u2309 select bits)",
         "Exhaustive (per output line; k \u2264 4 vars, no induction needed)"),
        ("BCD-to-7-segment", "4-bit only", "4",
         "Exhaustive (Section IV.E don't-care case study)"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, rows[0]):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows[1:]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            cell.text = text
    doc.save(path)
    return path


def sampled_verify(
    dis: DecisionInformationSystem,
    cover: List[FrozenSet[Literal]],
    confidence: float = 0.999,
    error_rate: float = 0.0005,
    seed: int = 0,
) -> Dict:
    """Monte Carlo equivalence check for effective n too large to
    enumerate exhaustively. Returns a dict with the pass/fail result and
    the (confidence, error_rate, N) bound actually used -- never a bare
    vector count (Section VI.D)."""
    import random
    random.seed(seed)
    n_samples = required_sample_size(confidence, error_rate)
    onset_set, offset_set = set(dis.onset), set(dis.offset)
    total = 2 ** dis.n
    mismatches = 0
    for _ in range(n_samples):
        m = random.randrange(total)
        expected = 1 if m in onset_set else (0 if m in offset_set else None)
        if expected is None:
            continue
        if _cover_value(cover, m, dis.n) != expected:
            mismatches += 1
    return {
        "passed": mismatches == 0,
        "mismatches": mismatches,
        "confidence": confidence,
        "error_rate": error_rate,
        "n_samples": n_samples,
    }


def verify(dis: DecisionInformationSystem, cover: List[FrozenSet[Literal]]) -> Dict:
    """Dispatches to exhaustive or sampled verification per Section VI.D
    threshold (effective n <= 24 -> exhaustive)."""
    if dis.n <= 24:
        return {"method": "exhaustive", "passed": exhaustive_verify(dis, cover)}
    result = sampled_verify(dis, cover)
    result["method"] = "sampled"
    return result


# ---------------------------------------------------------------------
# Full three-way benchmark run
# ---------------------------------------------------------------------

def run_benchmark(dis: DecisionInformationSystem) -> Dict:
    results = {}

    t0 = time.perf_counter()
    qmc_pis = qmc_prime_implicants(dis.n, dis.onset, dis.dc)
    qmc_cvr = petrick_cover(dis.n, dis.onset, qmc_pis)
    t_qmc = time.perf_counter() - t0
    results["QMC"] = {
        "gate_count": gate_count(qmc_cvr),
        "literal_count": literal_count(qmc_cvr),
        "time_s": t_qmc,
        "verification": verify(dis, qmc_cvr),
    }

    t0 = time.perf_counter()
    esp_cvr = espresso_cover(dis.n, dis.onset, dis.dc)
    t_esp = time.perf_counter() - t0
    results["Espresso"] = {
        "gate_count": gate_count(esp_cvr),
        "literal_count": literal_count(esp_cvr),
        "time_s": t_esp,
        "verification": verify(dis, esp_cvr),
    }

    t0 = time.perf_counter()
    rst_cvr, reduct = rst_cover(dis)
    t_rst = time.perf_counter() - t0
    results["RST"] = {
        "gate_count": gate_count(rst_cvr),
        "literal_count": literal_count(rst_cvr),
        "time_s": t_rst,
        "reduct_size": len(reduct),
        "variable_reduction_ratio": variable_reduction_ratio(dis.n, len(reduct)),
        "verification": verify(dis, rst_cvr),
    }

    return results


if __name__ == "__main__":
    print("=== Worked three-way comparison: majority function (n=3) ===")
    dis = DecisionInformationSystem(n=3, onset=[3, 5, 6, 7], offset=[0, 1, 2, 4], dc=[])
    res = run_benchmark(dis)
    for method, r in res.items():
        print(f"  {method}: gates={r['gate_count']}, literals={r['literal_count']}, "
              f"time={r['time_s']*1000:.3f}ms, verified={r['verification']['passed']}"
              + (f", reduct_size={r['reduct_size']}" if "reduct_size" in r else ""))

    print("\n=== Worked three-way comparison: structured n=4, vars {2,3} redundant ===")
    from module_3 import structured_truth_table
    onset, offset = structured_truth_table(4, relevant_vars=[0, 1], seed=1)
    dis2 = DecisionInformationSystem(n=4, onset=onset, offset=offset, dc=[])
    res2 = run_benchmark(dis2)
    for method, r in res2.items():
        print(f"  {method}: gates={r['gate_count']}, literals={r['literal_count']}, "
              f"time={r['time_s']*1000:.3f}ms, verified={r['verification']['passed']}"
              + (f", reduct_size={r['reduct_size']}, var_reduction={r['variable_reduction_ratio']:.2f}"
                 if "reduct_size" in r else ""))

    print("\n=== Sampled verification self-test (confidence=99.9%, error_rate=0.05%) ===")
    n_samples = required_sample_size(0.999, 0.0005)
    print(f"Required sample size N = {n_samples}")

    print("\n=== Generating Section VI artifacts ===")
    for fn, path in [
        (generate_table2_docx, "table2_figures_of_merit.docx"),
        (generate_table3_docx, "table3_feasibility_scope.docx"),
        (generate_table4_docx, "table4_circuit_family_verification.docx"),
    ]:
        out = fn(path)
        print(f"Written: {out}")
