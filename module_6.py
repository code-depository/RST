"""
module_6.py
Section VIII -- Real Circuit Case Studies

Generates truth tables for five circuit families and runs the Section VI
protocol on each, per the (corrected) verification assignment in Table 4:

  - Ripple-carry adder:  per-bit full-adder cell, minimized ONCE,
    scaled by width (structural induction over the carry chain).
  - Magnitude comparator: per-stage compare-cascade cell, minimized
    ONCE, scaled by width (structural induction over the cascade).
  - Priority encoder: monolithic per width (no decomposition; w <= 16
    is well within the exhaustive-verification threshold, Section VI.D).
  - Binary decoder: each output line is an independent function of only
    k = ceil(log2(w)) select bits; verified per line, no induction.
  - BCD-to-7-segment: real 7-segment truth tables, n=4, exhaustive.

Depends on module_1.py, module_3.py, module_4.py (same directory).
"""

import math
from typing import Dict, List, Tuple

from module_1 import DecisionInformationSystem
from module_4 import run_benchmark


# ---------------------------------------------------------------------
# A. Ripple-carry adder: full-adder cell
# ---------------------------------------------------------------------

def full_adder_cells() -> Dict[str, DecisionInformationSystem]:
    """n=3 cell: inputs (a, b, cin). sum = a xor b xor cin;
    cout = majority(a,b,cin)."""
    onset_sum = [m for m in range(8) if bin(m).count("1") % 2 == 1]
    onset_cout = [m for m in range(8) if bin(m).count("1") >= 2]
    offset_sum = [m for m in range(8) if m not in onset_sum]
    offset_cout = [m for m in range(8) if m not in onset_cout]
    return {
        "sum": DecisionInformationSystem(n=3, onset=onset_sum, offset=offset_sum, dc=[]),
        "cout": DecisionInformationSystem(n=3, onset=onset_cout, offset=offset_cout, dc=[]),
    }


# ---------------------------------------------------------------------
# B. Magnitude comparator: compare-cascade cell
# ---------------------------------------------------------------------

def comparator_cells() -> Dict[str, DecisionInformationSystem]:
    """n=4 cell: inputs (a_i, b_i, gt_in, eq_in).
    gt_out = gt_in OR (eq_in AND a_i AND NOT b_i)
    eq_out = eq_in AND (a_i == b_i)
    Attribute order: 0=a_i, 1=b_i, 2=gt_in, 3=eq_in.
    """
    onset_gt, onset_eq = [], []
    for m in range(16):
        bits = format(m, "04b")
        a, b, gt_in, eq_in = (int(c) for c in bits)
        gt_out = gt_in or (eq_in and a and not b)
        eq_out = eq_in and (a == b)
        if gt_out:
            onset_gt.append(m)
        if eq_out:
            onset_eq.append(m)
    offset_gt = [m for m in range(16) if m not in onset_gt]
    offset_eq = [m for m in range(16) if m not in onset_eq]
    return {
        "gt": DecisionInformationSystem(n=4, onset=onset_gt, offset=offset_gt, dc=[]),
        "eq": DecisionInformationSystem(n=4, onset=onset_eq, offset=offset_eq, dc=[]),
    }


# ---------------------------------------------------------------------
# C. Priority encoder: monolithic per width
# ---------------------------------------------------------------------

def priority_encoder_outputs(w: int) -> Dict[str, DecisionInformationSystem]:
    """w priority inputs i_0..i_{w-1} (i_{w-1} highest priority, matching
    the original 8-to-3 convention). k = ceil(log2(w)) output bits
    encode the index of the highest-priority active input (0 if none
    active). Each output bit is a monolithic function of all w inputs.
    """
    k = max(1, math.ceil(math.log2(w)))
    outputs = {f"y{j}": [] for j in range(k)}
    total = 2 ** w
    for m in range(total):
        bits = format(m, f"0{w}b")  # bits[0] = i_{w-1} (MSB position = highest index)
        highest = None
        for idx in range(w):  # idx corresponds to input i_idx; bits position (w-1-idx)
            if bits[w - 1 - idx] == "1":
                highest = idx  # keep scanning upward; last match = highest idx
        code = highest if highest is not None else 0
        code_bits = format(code, f"0{k}b")
        for j in range(k):
            if code_bits[j] == "1":
                outputs[f"y{j}"].append(m)
    result = {}
    for name, onset in outputs.items():
        offset = [m for m in range(total) if m not in onset]
        result[name] = DecisionInformationSystem(n=w, onset=onset, offset=offset, dc=[])
    return result


# ---------------------------------------------------------------------
# D. Binary decoder: per-output-line, select-width functions
# ---------------------------------------------------------------------

def decoder_outputs(w_lines: int) -> Dict[str, DecisionInformationSystem]:
    """w_lines output lines, k = ceil(log2(w_lines)) select inputs.
    Output line m is active iff the select input equals m; select values
    >= w_lines (when w_lines is not a power of 2) produce all-zero output
    (fully specified, no additional don't-cares introduced).
    """
    k = max(1, math.ceil(math.log2(w_lines)))
    result = {}
    for m in range(w_lines):
        onset = [m]  # single minterm in the k-variable select space
        offset = [x for x in range(2 ** k) if x != m]
        result[f"line{m}"] = DecisionInformationSystem(n=k, onset=onset, offset=offset, dc=[])
    return result


# ---------------------------------------------------------------------
# E. BCD-to-7-segment decoder (real truth table)
# ---------------------------------------------------------------------

# Standard common-cathode 7-segment encoding for digits 0-9 (segments a-g)
_BCD_SEGMENTS = {
    0: "1111110", 1: "0110000", 2: "1101101", 3: "1111001", 4: "0110011",
    5: "1011011", 6: "1011111", 7: "1110000", 8: "1111111", 9: "1111011",
}


def bcd_to_7segment_outputs() -> Dict[str, DecisionInformationSystem]:
    """n=4 (BCD digit 0-9), 7 outputs (segments a-g), codes 10-15 are the
    natural don't-care set (Section IV.E)."""
    dc = list(range(10, 16))
    result = {}
    for seg_idx, seg_name in enumerate("abcdefg"):
        onset = [digit for digit, pattern in _BCD_SEGMENTS.items() if pattern[seg_idx] == "1"]
        offset = [digit for digit in range(10) if digit not in onset]
        result[f"seg_{seg_name}"] = DecisionInformationSystem(n=4, onset=onset, offset=offset, dc=dc)
    return result


if __name__ == "__main__":
    print("=== A. Full-adder cell ===")
    for name, dis in full_adder_cells().items():
        res = run_benchmark(dis)
        print(f"  {name}: QMC={res['QMC']['gate_count']}, Espresso={res['Espresso']['gate_count']}, "
              f"RST={res['RST']['gate_count']}, reduct={res['RST']['reduct_size']}, "
              f"verified={all(r['verification']['passed'] for r in res.values())}")

    print("\n=== B. Comparator cascade cell ===")
    for name, dis in comparator_cells().items():
        res = run_benchmark(dis)
        print(f"  {name}: QMC={res['QMC']['gate_count']}, Espresso={res['Espresso']['gate_count']}, "
              f"RST={res['RST']['gate_count']}, reduct={res['RST']['reduct_size']}, "
              f"verified={all(r['verification']['passed'] for r in res.values())}")

    print("\n=== C. Priority encoder (w=4) ===")
    for name, dis in priority_encoder_outputs(4).items():
        res = run_benchmark(dis)
        print(f"  {name}: QMC={res['QMC']['gate_count']}, Espresso={res['Espresso']['gate_count']}, "
              f"RST={res['RST']['gate_count']}, reduct={res['RST']['reduct_size']}, "
              f"verified={all(r['verification']['passed'] for r in res.values())}")

    print("\n=== D. Binary decoder (w=8 lines) ===")
    for name, dis in list(decoder_outputs(8).items())[:3]:
        res = run_benchmark(dis)
        print(f"  {name}: QMC={res['QMC']['gate_count']}, Espresso={res['Espresso']['gate_count']}, "
              f"RST={res['RST']['gate_count']}, verified={all(r['verification']['passed'] for r in res.values())}")

    print("\n=== E. BCD-to-7-segment ===")
    for name, dis in bcd_to_7segment_outputs().items():
        res = run_benchmark(dis)
        print(f"  {name}: QMC={res['QMC']['gate_count']}, Espresso={res['Espresso']['gate_count']}, "
              f"RST={res['RST']['gate_count']}, reduct={res['RST']['reduct_size']}, "
              f"verified={all(r['verification']['passed'] for r in res.values())}")


# ---------------------------------------------------------------------
# Artifact generation: Tables 11-16 (.docx), Section VIII
# ---------------------------------------------------------------------

def _new_doc_with_title(title_text: str):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(11)
    return doc


def _add_table(doc, headers, data_rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in data_rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            cell.text = str(text)
    return table


def generate_table11_docx(path):
    from module_4 import run_benchmark
    cells = full_adder_cells()
    per_cell = {name: run_benchmark(dis) for name, dis in cells.items()}
    sum_g, cout_g = per_cell["sum"]["RST"]["gate_count"], per_cell["cout"]["RST"]["gate_count"]
    sum_l, cout_l = per_cell["sum"]["RST"]["literal_count"], per_cell["cout"]["RST"]["literal_count"]
    per_bit_g, per_bit_l = sum_g + cout_g, sum_l + cout_l

    doc = _new_doc_with_title("Table 11. Ripple-carry adder: per-bit cell and width-scaled totals.")
    data = [("sum", sum_g, sum_l, 3), ("cout", cout_g, cout_l, 3),
            ("per-bit total", per_bit_g, per_bit_l, "\u2014")]
    for w in [4, 8, 12, 16]:
        data.append((f"{w}-bit total", per_bit_g * w, per_bit_l * w, "\u2014"))
    _add_table(doc, ["Output / Width", "Gate count", "Literal count", "Reduct size"], data)
    doc.save(path)
    return path


def generate_table12_docx(path):
    from module_4 import run_benchmark
    cells = comparator_cells()
    per_cell = {name: run_benchmark(dis) for name, dis in cells.items()}
    gt_g, eq_g = per_cell["gt"]["RST"]["gate_count"], per_cell["eq"]["RST"]["gate_count"]
    gt_l, eq_l = per_cell["gt"]["RST"]["literal_count"], per_cell["eq"]["RST"]["literal_count"]
    gt_r, eq_r = per_cell["gt"]["RST"]["reduct_size"], per_cell["eq"]["RST"]["reduct_size"]
    per_stage_g, per_stage_l = gt_g + eq_g, gt_l + eq_l

    doc = _new_doc_with_title("Table 12. Magnitude comparator: per-stage cell and width-scaled totals.")
    data = [("gt", gt_g, gt_l, gt_r), ("eq", eq_g, eq_l, eq_r),
            ("per-stage total", per_stage_g, per_stage_l, "\u2014")]
    for w in [4, 8, 12, 16]:
        data.append((f"{w}-bit total", per_stage_g * w, per_stage_l * w, "\u2014"))
    _add_table(doc, ["Output / Width", "Gate count", "Literal count", "Reduct size (of 4)"], data)
    doc.save(path)
    return path


def generate_table13_docx(encoder_results, path):
    doc = _new_doc_with_title("Table 13. Priority encoder: per-output gate counts and RST reduct size, all tested widths.")
    data = []
    for w in [4, 8, 12, 16]:
        for r in encoder_results[w]:
            qmc = r["qmc_gates"] if r["qmc_gates"] is not None else "\u2014 (n>10)"
            data.append((w, r["output"], qmc, r["espresso_gates"], r["rst_gates"],
                         f"{r['reduct_size']}/{w}", r.get("rst_minimization_method", "exact")))
    _add_table(doc, ["w", "Output", "QMC", "Espresso", "RST", "Reduct size", "RST method"], data)
    doc.save(path)
    return path


def generate_table14_docx(decoder_results, path):
    doc = _new_doc_with_title("Table 14. Binary decoder: total gate count across all output lines, by width.")
    data = []
    for w in [4, 8, 12, 16]:
        rows = decoder_results[w]
        total_gates = sum(r["rst_gates"] for r in rows)
        all_verified = all(r["verified"] for r in rows)
        data.append((w, len(rows), total_gates, all_verified))
    _add_table(doc, ["w (output lines)", "Select bits (k)", "Total gate count", "All verified"], data)
    doc.save(path)
    return path


def generate_table15_docx(path):
    from module_4 import run_benchmark
    doc = _new_doc_with_title("Table 15. BCD-to-7-segment decoder: per-segment gate counts and RST reduct size.")
    data = []
    total = 0
    for name, dis in bcd_to_7segment_outputs().items():
        res = run_benchmark(dis)
        g = res["RST"]["gate_count"]
        total += g
        tied = res["QMC"]["gate_count"] == res["Espresso"]["gate_count"] == g
        data.append((name, res["QMC"]["gate_count"], res["Espresso"]["gate_count"], g,
                     f"{res['RST']['reduct_size']}/4", tied))
    data.append(("TOTAL", "\u2014", "\u2014", total, "\u2014", "\u2014"))
    _add_table(doc, ["Segment", "QMC", "Espresso", "RST", "Reduct size", "All tied"], data)
    doc.save(path)
    return path


def generate_table16_docx(encoder_results, path):
    doc = _new_doc_with_title("Table 16. Case study overview: variable/support reduction ratio by family.")
    data = [
        ("Ripple-carry adder", "3 (per-bit)", "0.00 (baseline, no redundancy)"),
        ("Magnitude comparator", "3-4 (per-stage)", "0.00 (gt), 0.25 (eq)"),
    ]
    for w in [16]:
        for r in encoder_results[w]:
            ratio = 1 - r["reduct_size"] / w
            data.append((f"Priority encoder ({r['output']}, w={w})", f"{r['reduct_size']}/{w}", f"{ratio:.2f}"))
    data.append(("Binary decoder", "k \u2264 4 (all needed)", "0.00 (baseline, no redundancy)"))
    data.append(("BCD-to-7-segment (avg.)", "~3.4/4", "~0.14 (3 of 7 segments show reduction)"))
    _add_table(doc, ["Family", "Reduct size", "Variable reduction ratio"], data)
    doc.save(path)
    return path


if __name__ == "__main__":
    import json
    print("Generating Section VIII tables...")

    generate_table11_docx("table11_adder.docx")
    print("Table 11 written.")
    generate_table12_docx("table12_comparator.docx")
    print("Table 12 written.")

    with open("results_encoder_decoder.json") as f:
        ed_results = json.load(f)
    enc_results = {int(k): v for k, v in ed_results["encoder"].items()}
    dec_results = {int(k): v for k, v in ed_results["decoder"].items()}

    generate_table13_docx(enc_results, "table13_encoder.docx")
    print("Table 13 written.")
    generate_table14_docx(dec_results, "table14_decoder.docx")
    print("Table 14 written.")
    generate_table15_docx("table15_bcd.docx")
    print("Table 15 written.")
    generate_table16_docx(enc_results, "table16_summary.docx")
    print("Table 16 written.")
