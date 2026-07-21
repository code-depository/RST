"""
module_1.py
Section III -- RST <-> QMC Formal Equivalence (Proposition 1)

Self-contained module: Decision Information System, literal-level
discernibility (Section III.A), CNF->DNF consensus conversion
(Section III.C, Step 1), a reference Quine-McCluskey implementation
(Section III.C, Step 3), and the equivalence check (Proposition 1).

Also generates the manuscript artifact for Section III, derived directly
from the module's own computed values (not hand-typed):
    - Table 1  (.docx): literal discernibility clauses for Delta_7,
      the worked majority-function example, cited inline in the proof
      (Section III.C, Step 1).

External dependencies (only for artifact generation, not for the core
proof/verification logic above): python-docx.
Run this file directly to execute the worked example, the randomized
equivalence sweep, and generate the artifact.
"""

import random
from dataclasses import dataclass, field
from typing import Dict, FrozenSet, Iterable, List, Optional, Set, Tuple

Literal = Tuple[int, int]  # (attribute_index, value)


# ---------------------------------------------------------------------
# Decision Information System  (Section II.B)
# ---------------------------------------------------------------------

@dataclass
class DecisionInformationSystem:
    n: int
    onset: List[int]
    offset: List[int]
    dc: List[int] = field(default_factory=list)

    def __post_init__(self):
        total = 2 ** self.n
        covered = set(self.onset) | set(self.offset) | set(self.dc)
        if len(covered) != total:
            raise ValueError(
                f"onset/offset/dc must partition all {total} minterms; "
                f"got {len(covered)} distinct minterms."
            )

    def bits(self, m: int) -> str:
        return format(m, f"0{self.n}b")


# ---------------------------------------------------------------------
# Literal-level discernibility  (Section III.A)
# ---------------------------------------------------------------------

def literal_discernibility_clauses(
    dis: DecisionInformationSystem, mi: int
) -> List[FrozenSet[Literal]]:
    """CNF clauses of Delta_i: one clause per offset minterm mj, each the
    set of literals (fixed to mi's values) that distinguish mi from mj."""
    bi = dis.bits(mi)
    clauses = []
    for mj in dis.offset:
        bj = dis.bits(mj)
        clause = frozenset((a, int(bi[a])) for a in range(dis.n) if bi[a] != bj[a])
        clauses.append(clause)
    return clauses


# ---------------------------------------------------------------------
# CNF -> DNF via distribution + absorption  (Section III.C, Step 1)
# ---------------------------------------------------------------------

def _literal_conflict(term: FrozenSet[Literal], lit: Literal) -> bool:
    attr, val = lit
    return any(a == attr and v != val for (a, v) in term)


def _absorb(terms: Iterable[FrozenSet[Literal]]) -> Set[FrozenSet[Literal]]:
    ordered = sorted(terms, key=len)
    kept: List[FrozenSet[Literal]] = []
    for t in ordered:
        if not any(k <= t for k in kept):
            kept.append(t)
    return set(kept)


def cnf_to_dnf(clauses: List[FrozenSet[Literal]]) -> List[FrozenSet[Literal]]:
    """Distribute clauses one at a time, absorbing after each step, to
    obtain the minimal (prime) DNF terms satisfying the CNF."""
    if not clauses:
        return [frozenset()]
    terms: Set[FrozenSet[Literal]] = {frozenset([lit]) for lit in clauses[0]}
    for clause in clauses[1:]:
        new_terms: Set[FrozenSet[Literal]] = set()
        for term in terms:
            for lit in clause:
                if lit in term:
                    new_terms.add(term)
                    continue
                if _literal_conflict(term, lit):
                    continue
                new_terms.add(frozenset(term | {lit}))
        terms = _absorb(new_terms)
    return list(terms)




    """Union over all onset minterms of Delta_i's DNF terms, globally
    absorbed (Section III.C, Step 2)."""
    all_terms: Set[FrozenSet[Literal]] = set()
    for mi in dis.onset:
        clauses = literal_discernibility_clauses(dis, mi)
        if not clauses:
            all_terms.add(frozenset())
            continue
        all_terms.update(cnf_to_dnf(clauses))
    return list(_absorb(all_terms))


def rst_prime_implicants(dis: DecisionInformationSystem) -> List[FrozenSet[Literal]]:
    """Union over all onset minterms of Delta_i's DNF terms, globally
    absorbed (Section III.C, Step 2)."""
    all_terms: Set[FrozenSet[Literal]] = set()
    for mi in dis.onset:
        clauses = literal_discernibility_clauses(dis, mi)
        if not clauses:
            all_terms.add(frozenset())
            continue
        all_terms.update(cnf_to_dnf(clauses))
    return list(_absorb(all_terms))


# ---------------------------------------------------------------------
# Reference Quine-McCluskey implementation  (Section III.C, Step 3)
# ---------------------------------------------------------------------

def _pattern_covers(pattern: str, m: int, n: int) -> bool:
    b = format(m, f"0{n}b")
    return all(pattern[i] in ("-", b[i]) for i in range(n))


def _pattern_to_literals(pattern: str) -> FrozenSet[Literal]:
    return frozenset((i, int(c)) for i, c in enumerate(pattern) if c != "-")


def _intmask_to_pattern(value: int, mask: int, n: int) -> str:
    chars = []
    for i in range(n - 1, -1, -1):
        bit = 1 << i
        if mask & bit:
            chars.append("-")
        else:
            chars.append("1" if value & bit else "0")
    return "".join(chars)


def qmc_prime_implicants(
    n: int, onset: List[int], dc: Optional[List[int]] = None
) -> List[FrozenSet[Literal]]:
    """Compute the complete set of prime implicants of f via QMC combining.
    dc minterms participate in combining (so implicants may use them) but
    a prime implicant is only retained if it covers >= 1 true onset minterm.

    Combining is done via integer (value, mask) pairs rather than
    character-by-character string comparison (the original implementation)
    -- the latter does not scale past a few hundred minterms in pure
    Python (measured directly: ~73s at 1365 minterms with the string
    version, vs a fraction of a second with this version at the same
    scale). value has a 0 at every masked ('-') position; two terms
    combine iff their masks match and their values differ in exactly
    one unmasked bit.
    """
    dc = dc or []
    all_terms = sorted(set(onset) | set(dc))
    if not all_terms:
        return []

    current: Set[Tuple[int, int]] = {(t, 0) for t in all_terms}
    primes: Set[Tuple[int, int]] = set()
    full_mask = (1 << n) - 1

    while True:
        used: Set[Tuple[int, int]] = set()
        next_gen: Set[Tuple[int, int]] = set()
        by_ones: Dict[Tuple[int, int], List[Tuple[int, int]]] = {}
        for (v, m) in current:
            ones = bin(v & ~m & full_mask).count("1")
            by_ones.setdefault((ones, m), []).append((v, m))

        for (ones, m), group in by_ones.items():
            other_group = by_ones.get((ones + 1, m))
            if not other_group:
                continue
            for (va, ma) in group:
                for (vb, mb) in other_group:
                    diff = va ^ vb
                    if diff != 0 and (diff & (diff - 1)) == 0:
                        next_gen.add((va & ~diff, ma | diff))
                        used.add((va, ma))
                        used.add((vb, mb))

        for term in current:
            if term not in used:
                primes.add(term)
        if not next_gen:
            break
        current = next_gen

    result = []
    for (v, m) in primes:
        pattern = _intmask_to_pattern(v, m, n)
        if any(_pattern_covers(pattern, mt, n) for mt in onset):
            result.append(_pattern_to_literals(pattern))
    return result


# ---------------------------------------------------------------------
# Proposition 1 equivalence check
# ---------------------------------------------------------------------

def check_proposition1(dis: DecisionInformationSystem):
    """Returns (equal: bool, rst_only: set, qmc_only: set)."""
    rst_pis = set(rst_prime_implicants(dis))
    qmc_pis = set(qmc_prime_implicants(dis.n, dis.onset, dis.dc))
    return rst_pis == qmc_pis, rst_pis - qmc_pis, qmc_pis - rst_pis


def to_expression(implicants: List[FrozenSet[Literal]]) -> str:
    if not implicants:
        return "0"
    terms = []
    for pi in sorted(implicants, key=lambda s: sorted(s)):
        if not pi:
            return "1"
        term = "".join(f"x{a}" if v else f"x{a}'" for a, v in sorted(pi))
        terms.append(term)
    return " + ".join(terms)


def random_truth_table(n: int, dc_fraction: float = 0.0, seed: Optional[int] = None):
    if seed is not None:
        random.seed(seed)
    total = 2 ** n
    minterms = list(range(total))
    random.shuffle(minterms)
    n_dc = int(round(dc_fraction * total))
    dc = sorted(minterms[:n_dc])
    remaining = minterms[n_dc:]
    split = len(remaining) // 2
    return sorted(remaining[:split]), sorted(remaining[split:]), dc


# ---------------------------------------------------------------------
# Artifact generation: Table 1 (.docx)
# ---------------------------------------------------------------------

def _lit_str(lit: Literal) -> str:
    a, v = lit
    return f"x{a}" if v else f"x{a}'"


def _clause_str(clause: FrozenSet[Literal]) -> str:
    return " + ".join(_lit_str(l) for l in sorted(clause)) if clause else "(empty)"


def _term_str(term: FrozenSet[Literal]) -> str:
    return "".join(_lit_str(l) for l in sorted(term)) if term else "1"


def generate_table1_docx(dis: DecisionInformationSystem, mi: int, path: str) -> str:
    """Table 1: literal discernibility clauses of Delta_mi (Section III.B),
    values computed directly from literal_discernibility_clauses -- not
    hand-entered.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    clauses = literal_discernibility_clauses(dis, mi)

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(
        f"Table 1. Literal discernibility clauses of \u0394\u2087 "
        f"for onset minterm m\u1D62 = {mi} ({dis.bits(mi)}), "
        f"majority function f(x0,x1,x2) = x0x1 + x0x2 + x1x2."
    )
    run.bold = True
    run.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    headers = ["m\u2c7c (offset)", "Bits", "Discerning attributes", "Clause literals (fixed to m\u1D62)"]
    for cell, text in zip(hdr, headers):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for mj, clause in zip(dis.offset, clauses):
        row = table.add_row().cells
        attrs = sorted({a for a, _ in clause})
        row[0].text = str(mj)
        row[1].text = dis.bits(mj)
        row[2].text = ", ".join(f"x{a}" for a in attrs)
        row[3].text = _clause_str(clause)

    doc.save(path)
    return path




if __name__ == "__main__":
    print("=== Worked example: majority function f = x0x1 + x0x2 + x1x2 ===")
    dis = DecisionInformationSystem(n=3, onset=[3, 5, 6, 7], offset=[0, 1, 2, 4], dc=[])
    equal, rst_only, qmc_only = check_proposition1(dis)
    print("Proposition 1 holds:", equal)
    print("RST PIs:", to_expression(rst_prime_implicants(dis)))
    print("QMC PIs:", to_expression(qmc_prime_implicants(3, [3, 5, 6, 7], [])))

    print("\n=== Randomized sweep ===")
    n_values, dc_fractions, instances = [3, 4, 5, 6], [0.0, 0.1, 0.2, 0.3], 10
    total_checked, total_passed = 0, 0
    for n in n_values:
        for dcf in dc_fractions:
            for seed in range(instances):
                onset, offset, dc = random_truth_table(n, dcf, seed=seed * 97 + n)
                if not onset or not offset:
                    continue
                d = DecisionInformationSystem(n=n, onset=onset, offset=offset, dc=dc)
                eq, _, _ = check_proposition1(d)
                total_checked += 1
                total_passed += int(eq)
    print(f"Proposition 1 confirmed on {total_passed}/{total_checked} random instances "
          f"(n in {n_values}, DC in {dc_fractions}).")

    print("\n=== Generating Section III artifact ===")
    table_path = generate_table1_docx(dis, mi=7, path="table1_discernibility_clauses.docx")
    print(f"Table 1 written to: {table_path}")
